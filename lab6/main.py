import toga
from toga.style import Pack
from toga.style.pack import COLUMN, ROW
from appliances.iron import Iron
from appliances.tv import TV
from appliances.washing_machine import WashingMachine
from docx import Document
import openpyxl

class ApplianceApp(toga.App):
    def startup(self):
        main_box = toga.Box(style=Pack(direction=COLUMN, padding=10))

        # Input fields
        self.power_input = toga.TextInput(placeholder='Power (W)', style=Pack(flex=1))
        self.usage_input = toga.TextInput(placeholder='Usage per day (hours)', style=Pack(flex=1))
        self.days_input = toga.TextInput(placeholder='Days', style=Pack(flex=1))
        self.cost_input = toga.TextInput(placeholder='Cost per kWh', style=Pack(flex=1))

        # Dropdown to select appliance
        self.appliance_select = toga.Selection(items=['Iron', 'TV', 'Washing Machine'], style=Pack(flex=1))

        # Button to calculate
        calculate_button = toga.Button('Calculate', on_press=self.calculate, style=Pack(padding=5))

        # Result display
        self.result_label = toga.Label('Result: ', style=Pack(padding=5))

        input_box = toga.Box(style=Pack(direction=ROW, padding=5))
        input_box.add(self.power_input)
        input_box.add(self.usage_input)
        input_box.add(self.days_input)
        input_box.add(self.cost_input)
        input_box.add(self.appliance_select)
        input_box.add(calculate_button)

        main_box.add(input_box)
        main_box.add(self.result_label)

        # Save buttons
        save_docx_button = toga.Button('Save as DOCX', on_press=self.save_docx, style=Pack(padding=5))
        save_xlsx_button = toga.Button('Save as XLSX', on_press=self.save_xlsx, style=Pack(padding=5))

        main_box.add(save_docx_button)
        main_box.add(save_xlsx_button)

        self.main_window = toga.MainWindow(title=self.formal_name)
        self.main_window.content = main_box
        self.main_window.show()

    def calculate(self, widget):
        power = float(self.power_input.value)
        usage_per_day = float(self.usage_input.value)
        days = int(self.days_input.value)
        cost_per_kwh = float(self.cost_input.value)
        appliance_type = self.appliance_select.value

        if appliance_type == 'Iron':
            appliance = Iron(power, usage_per_day)
        elif appliance_type == 'TV':
            appliance = TV(power, usage_per_day)
        elif appliance_type == 'Washing Machine':
            appliance = WashingMachine(power, usage_per_day)

        energy_consumption = appliance.calculate_energy_consumption(days)
        cost = appliance.calculate_cost(days, cost_per_kwh)

        self.result_label.text = f'Result: {energy_consumption:.2f} kWh, ${cost:.2f}'

    def save_docx(self, widget):
        doc = Document()
        doc.add_heading('Appliance Energy Consumption Report', 0)
        doc.add_paragraph(self.result_label.text)
        doc.save('report.docx')

    def save_xlsx(self, widget):
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = 'Report'
        ws['A1'] = 'Appliance Energy Consumption Report'
        ws['A2'] = self.result_label.text
        wb.save('report.xlsx')

def main():
    return ApplianceApp('Appliance Energy Consumption', 'org.beeware.appliance')

if __name__ == '__main__':
    main().main_loop()