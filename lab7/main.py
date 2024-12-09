import tkinter as tk
from tkinter import messagebox
from docx import Document
import openpyxl
from Appliances.Iron import Iron
from Appliances.TV import TV
from Appliances.WashingMachine import WashingMachine
from Appliances.Appliance import Appliance


class ApplianceApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Appliances Energy Consumption")

        # Create widgets
        self.create_widgets()

    def create_widgets(self):
        # Input fields
        self.power_label = tk.Label(self.root, text="Power (W):")
        self.power_label.grid(row=0, column=0, padx=10, pady=5)
        self.power_input = tk.Entry(self.root)
        self.power_input.grid(row=0, column=1, padx=10, pady=5)

        self.usage_label = tk.Label(self.root, text="Usage per day (hours):")
        self.usage_label.grid(row=1, column=0, padx=10, pady=5)
        self.usage_input = tk.Entry(self.root)
        self.usage_input.grid(row=1, column=1, padx=10, pady=5)

        self.days_label = tk.Label(self.root, text="Days:")
        self.days_label.grid(row=2, column=0, padx=10, pady=5)
        self.days_input = tk.Entry(self.root)
        self.days_input.grid(row=2, column=1, padx=10, pady=5)

        self.cost_label = tk.Label(self.root, text="Cost per kWh:")
        self.cost_label.grid(row=3, column=0, padx=10, pady=5)
        self.cost_input = tk.Entry(self.root)
        self.cost_input.grid(row=3, column=1, padx=10, pady=5)

        # Dropdown to select appliance
        self.appliance_label = tk.Label(self.root, text="Select Appliances:")
        self.appliance_label.grid(row=4, column=0, padx=10, pady=5)
        self.appliance_var = tk.StringVar()
        self.appliance_select = tk.OptionMenu(self.root, self.appliance_var, 'Iron', 'TV', 'Washing Machine')
        self.appliance_select.grid(row=4, column=1, padx=10, pady=5)
        self.appliance_var.set('Iron')  # Default value

        # Button to calculate
        self.calculate_button = tk.Button(self.root, text="Calculate", command=self.calculate)
        self.calculate_button.grid(row=5, column=0, columnspan=2, pady=10)

        # Result label
        self.result_label = tk.Label(self.root, text="Result: ")
        self.result_label.grid(row=6, column=0, columnspan=2, padx=10, pady=5)

        # Save buttons
        self.save_docx_button = tk.Button(self.root, text="Save as DOCX", command=self.save_docx)
        self.save_docx_button.grid(row=7, column=0, pady=10)
        self.save_xlsx_button = tk.Button(self.root, text="Save as XLSX", command=self.save_xlsx)
        self.save_xlsx_button.grid(row=7, column=1, pady=10)

    def calculate(self):
        try:
            # Get user input
            power = float(self.power_input.get())
            usage_per_day = float(self.usage_input.get())
            days = int(self.days_input.get())
            cost_per_kwh = float(self.cost_input.get())
            appliance_type = self.appliance_var.get()

            # Create appliance object based on selection
            if appliance_type == 'Iron':
                appliance = Iron(power, usage_per_day)
            elif appliance_type == 'TV':
                appliance = TV(power, usage_per_day)
            elif appliance_type == 'Washing Machine':
                appliance = WashingMachine(power, usage_per_day)

            # Calculate energy consumption and cost
            energy_consumption = appliance.calculate_energy_consumption(days)
            cost = appliance.calculate_cost(days, cost_per_kwh)

            # Display result
            self.result_label.config(text=f"Result: {energy_consumption:.2f} kWh, ${cost:.2f}")

        except ValueError:
            messagebox.showerror("Input Error", "Please enter valid values.")

    def save_docx(self):
        doc = Document()
        doc.add_heading('Appliances Energy Consumption Report', 0)
        doc.add_paragraph(self.result_label.cget("text"))
        doc.save('report.docx')

    def save_xlsx(self):
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = 'Report'
        ws['A1'] = 'Appliances Energy Consumption Report'
        ws['A2'] = self.result_label.cget("text")
        wb.save('report.xlsx')


def main():
    root = tk.Tk()
    app = ApplianceApp(root)
    root.mainloop()


if __name__ == '__main__':
    main()
